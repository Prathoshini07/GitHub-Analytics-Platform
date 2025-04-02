import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict

def add_page_number(paragraph):
    """Add page numbers in the footer safely"""
    if not paragraph:
        return  

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._element.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)

def categorize_issues(issues):
    """Categorize issues by type, severity, and component"""
    by_type = defaultdict(list)
    by_severity = defaultdict(list)
    by_component = defaultdict(list)
    
    for issue in issues:
        by_type[issue['type']].append(issue)
        by_severity[issue['severity']].append(issue)
        component = issue['component'].split(':', 1)[1] if ':' in issue['component'] else issue['component']
        by_component[component].append(issue)
    
    return by_type, by_severity, by_component

def get_severity_color(severity):
    """Return RGB color based on severity"""
    colors = {
        "BLOCKER": RGBColor(192, 0, 0), 
        "CRITICAL": RGBColor(255, 0, 0), 
        "MAJOR": RGBColor(255, 127, 0), 
        "MINOR": RGBColor(255, 201, 14), 
        "INFO": RGBColor(0, 0, 255) 
    }
    return colors.get(severity, RGBColor(0, 0, 0))  

def format_issue_table(table, issues):
    """Format the issue table with standardized styling"""
    headers = ["Issue Key", "Type", "Severity", "File", "Line", "Message", "Effort"]
    
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
    
    for row_idx, issue in enumerate(issues, start=1):
        row = table.add_row()
        component = issue['component'].split(':', 1)[1] if ':' in issue['component'] else issue['component']
        file_name = component.split('/')[-1] if '/' in component else component

        issue_data = [
            issue['key'], issue['type'], issue['severity'], file_name,
            str(issue.get('line', 'N/A')), issue['message'], issue['effort']
        ]
        
        for col_idx, data in enumerate(issue_data):
            cell = row.cells[col_idx]
            cell.text = data
            if col_idx == 2: 
                run = cell.paragraphs[0].runs[0]
                run.font.color.rgb = get_severity_color(issue['severity'])
                run.bold = True

def main():
    """Generate a SonarCloud report in a Word document"""
    with open("paste3.json", "r", encoding="utf-8") as file:
        data = json.load(file)
    
    issues = data.get('issues', [])
    if not issues:
        print("Warning: No issues found in JSON file.")
    
    doc = Document()
    doc.add_heading("SonarCloud Code Quality Report", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    project_name = issues[0].get('projectName', "Unknown Project") if issues else "Unknown Project"
    update_date = issues[0].get('updateDate', 'N/A').split('T')[0] if issues else 'N/A'

    doc.add_paragraph(f"Project: {project_name}", style="Heading 2").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on: {update_date}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    doc.add_heading("Executive Summary", level=1)

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.cell(0, 0).text = "Metric"
    summary_table.cell(0, 1).text = "Value"

    metrics = [
        ("Total Issues", str(data['total'])),
        ("Effort Required", f"{data['effortTotal']} mins"),
        ("Technical Debt", f"{data['debtTotal']} mins")
    ]
    
    by_type, by_severity, by_component = categorize_issues(issues)

    for severity in ["BLOCKER", "CRITICAL", "MAJOR", "MINOR", "INFO"]:
        if severity in by_severity:
            metrics.append((f"{severity} Issues", str(len(by_severity[severity]))))
    
    for issue_type in ["BUG", "VULNERABILITY", "CODE_SMELL"]:
        if issue_type in by_type:
            metrics.append((f"{issue_type.replace('_', ' ').title()}", str(len(by_type[issue_type]))))
    
    for metric, value in metrics:
        row = summary_table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = value

    doc.add_page_break()
    doc.add_heading("All Issues", level=1)

    all_issues_table = doc.add_table(rows=1, cols=7)
    all_issues_table.style = 'Table Grid'
    format_issue_table(all_issues_table, issues)

    section = doc.sections[0]
    footer = section.footer
    if not footer.paragraphs:
        footer.paragraphs.append(footer.add_paragraph())
    add_page_number(footer.paragraphs[0])

    doc.save("SonarCloud_Detailed_Report.docx")
    print("Document created: SonarCloud_Detailed_Report.docx")

if __name__ == "__main__":
    main()